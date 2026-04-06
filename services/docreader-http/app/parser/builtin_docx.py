"""
builtin_docx：python-docx 回退路径；抽取段落、表格，并尽量提取段落内嵌图片（参考 WeKnora Docx 内联图逻辑，单进程简化版）。
"""

from __future__ import annotations

import base64
import io
import logging
import uuid
from typing import Any, Optional

from docx import Document as DocxDocument
from docx.image.exceptions import (
    InvalidImageStreamError,
    UnexpectedEndOfFileError,
    UnrecognizedImageError,
)
from docx.opc.oxml import parse_xml
from docx.opc.pkgreader import _SerializedRelationship, _SerializedRelationships
from PIL import Image

from app.models.document import Document as DocModel
from app.parser.base_parser import BaseParser

logger = logging.getLogger(__name__)


def _apply_docx_opc_patch() -> None:
    """
    修复部分 DOCX 中异常 relationship 导致 python-docx 解析失败的问题（WeKnora 同源补丁）。
    """

    def load_from_xml_v2(baseURI, rels_item_xml):  # type: ignore[no-untyped-def]
        srels = _SerializedRelationships()
        if rels_item_xml is not None:
            rels_elm = parse_xml(rels_item_xml)
            for rel_elm in rels_elm.Relationship_lst:
                if rel_elm.target_ref in ("../NULL", "NULL"):
                    continue
                srels._srels.append(_SerializedRelationship(baseURI, rel_elm))
        return srels

    _SerializedRelationships.load_from_xml = load_from_xml_v2  # type: ignore[method-assign]


_apply_docx_opc_patch()


def _image_from_paragraph(document: DocxDocument, paragraph) -> Optional[bytes]:  # type: ignore[no-untyped-def]
    """从段落中提取第一张内嵌图片的原始字节（PNG 优先重新编码以统一格式）。"""
    imgs = paragraph._element.xpath(".//pic:pic")
    if not imgs:
        return None
    pic = imgs[0]
    try:
        embed = pic.xpath(".//a:blip/@r:embed")[0]
        related_part = document.part.related_parts[embed]
    except Exception as e:
        logger.debug("docx embed resolve skip: %s", e)
        return None
    try:
        blob = related_part.image.blob
    except (UnrecognizedImageError, UnexpectedEndOfFileError, InvalidImageStreamError) as e:
        logger.debug("docx image blob skip: %s", e)
        return None
    except Exception as e:
        logger.debug("docx image blob error: %s", e)
        return None
    try:
        im = Image.open(io.BytesIO(blob))
        if im.width < 30 or im.height < 30:
            return None
        buf = io.BytesIO()
        if im.mode not in ("RGB", "RGBA"):
            im = im.convert("RGBA")
        im.save(buf, format="PNG")
        return buf.getvalue()
    except Exception as e:
        logger.debug("docx pil encode skip: %s", e)
        return None


class DocxBuiltinParser(BaseParser):
    """DocxBuiltinParser 在 MarkItDown 无效时提供结构化文本 + 内联图。"""

    def parse_into_text(self, content: bytes) -> DocModel:
        meta: dict[str, Any] = {"parser": "python-docx"}
        parts: list[str] = []
        images: dict[str, str] = {}
        try:
            doc = DocxDocument(io.BytesIO(content))
        except Exception as e:
            logger.exception("docx open failed")
            return DocModel(metadata={**meta, "error": str(e)})

        for para in doc.paragraphs:
            t = para.text.strip()
            if t:
                parts.append(t)
            raw_png = _image_from_paragraph(doc, para)
            if raw_png:
                ref = f"images/{uuid.uuid4().hex}.png"
                images[ref] = base64.standard_b64encode(raw_png).decode("ascii")
                parts.append(f"![]({ref})")

        for table in doc.tables:
            rows_html = []
            for row in table.rows:
                cells = " | ".join(c.text.strip() for c in row.cells if c.text.strip())
                if cells:
                    rows_html.append(cells)
            if rows_html:
                parts.append("## Table\n\n" + "\n\n".join(rows_html))

        text = "\n\n".join(parts)
        meta["paragraphs"] = len(doc.paragraphs)
        if not text.strip() and not images:
            return DocModel(metadata=meta)
        return DocModel(content=text, images=images, metadata=meta)
